from __future__ import annotations

import posixpath
import zipfile
from pathlib import Path
from xml.etree import ElementTree

from docx import Document


MAX_ENTRIES = 2_000
MAX_EXPANDED_BYTES = 250 * 1024 * 1024
MAX_ENTRY_BYTES = 50 * 1024 * 1024
MAX_COMPRESSION_RATIO = 100


def _inspect_container(path: Path) -> None:
    try:
        with zipfile.ZipFile(path) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_ENTRIES:
                raise ValueError("unsafe_docx_source")
            expanded = 0
            seen_names: set[str] = set()
            for entry in entries:
                name = entry.filename.replace("\\", "/")
                if name.startswith("/") or name.startswith("../") or "/../" in name or posixpath.normpath(name) != name:
                    raise ValueError("unsafe_docx_source")
                if entry.file_size > MAX_ENTRY_BYTES:
                    raise ValueError("unsafe_docx_source")
                if entry.compress_size and entry.file_size / entry.compress_size > MAX_COMPRESSION_RATIO:
                    raise ValueError("unsafe_docx_source")
                folded_name = name.casefold()
                if folded_name in seen_names:
                    raise ValueError("unsafe_docx_source")
                seen_names.add(folded_name)
                expanded += entry.file_size
                if expanded > MAX_EXPANDED_BYTES:
                    raise ValueError("unsafe_docx_source")
                lowered = name.lower()
                if lowered.endswith(("vbaproject.bin", ".xls", ".xlsx", ".ppt", ".pptx")) or lowered.startswith("word/embeddings/"):
                    raise ValueError("unsafe_docx_source")
                if lowered.endswith(".rels"):
                    content = archive.read(entry)
                    try:
                        root = ElementTree.fromstring(content)
                    except ElementTree.ParseError as exc:
                        raise ValueError("unsafe_docx_source") from exc
                    for relationship in root:
                        target_mode = relationship.attrib.get("TargetMode", "").lower()
                        target = relationship.attrib.get("Target", "")
                        lowered_target = target.lower()
                        if target_mode == "external" or lowered_target.startswith(("http:", "https:", "file:")):
                            raise ValueError("unsafe_docx_source")
    except ValueError:
        raise
    except (OSError, zipfile.BadZipFile) as exc:
        raise ValueError("invalid_docx_source") from exc


def extract_docx_evidence(path: str | Path) -> dict[str, object]:
    source = Path(path)
    _inspect_container(source)
    try:
        document = Document(source)
    except Exception as exc:
        raise ValueError("invalid_docx_source") from exc
    sections = [{"width": round(section.page_width / 914400, 4), "height": round(section.page_height / 914400, 4), "orientation": str(section.orientation)} for section in document.sections]
    return {"paragraph_count": len(document.paragraphs), "table_count": len(document.tables), "section_count": len(sections), "sections": sections}

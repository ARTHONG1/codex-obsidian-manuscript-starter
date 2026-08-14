import importlib.util
import sys
import tempfile
import unittest
from pathlib import Path

from PIL import Image

ROOT = Path(__file__).parents[1]
SCRIPTS = ROOT / "plugins/obsidian-manuscript-publisher/skills/obsidian-manuscript-publisher/scripts"
sys.path.insert(0, str(SCRIPTS))


def load(name):
    spec = importlib.util.spec_from_file_location(name, SCRIPTS / f"{name}.py")
    module = importlib.util.module_from_spec(spec)
    sys.modules[name] = module
    spec.loader.exec_module(module)
    return module


extract_image = extract_pdf = extract_docx = None
try:
    extract_image = load("extract_image_template")
    extract_pdf = load("extract_pdf_template")
    extract_docx = load("extract_docx_template")
except FileNotFoundError:
    pass


class TemplateExtractorTests(unittest.TestCase):
    def setUp(self):
        self.assertIsNotNone(extract_image, "template extractors must exist")

    def test_image_evidence_is_bounded_and_deterministic(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "sample.png"
            Image.new("RGB", (120, 80), "white").save(path)
            first = extract_image.extract_image_evidence(path)
            second = extract_image.extract_image_evidence(path)
            self.assertEqual(first, second)
            self.assertEqual(first["width"], 120)
            self.assertNotIn(str(path), str(first))

    def test_image_extractor_rejects_non_allowlisted_format(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "sample.gif"
            Image.new("RGB", (12, 8), "white").save(path, format="GIF")
            with self.assertRaisesRegex(ValueError, "unsupported_image_format"):
                extract_image.extract_image_evidence(path)

    def test_pdf_evidence_reports_pages_without_source_text(self):
        try:
            from reportlab.pdfgen import canvas
        except ImportError:
            self.skipTest("reportlab unavailable")
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "sample.pdf"
            pdf = canvas.Canvas(str(path))
            pdf.drawString(20, 20, "sample")
            pdf.save()
            evidence = extract_pdf.extract_pdf_evidence(path)
            self.assertEqual(evidence["page_count"], 1)
            self.assertNotIn(str(path), str(evidence))

    def test_pdf_extractor_rejects_encrypted_input(self):
        from pypdf import PdfWriter
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "encrypted.pdf"
            writer = PdfWriter()
            writer.add_blank_page(width=100, height=100)
            writer.encrypt("secret")
            with path.open("wb") as stream:
                writer.write(stream)
            with self.assertRaisesRegex(ValueError, "unsafe_pdf_source"):
                extract_pdf.extract_pdf_evidence(path)

    def test_docx_evidence_reports_paragraph_and_table_shapes(self):
        from docx import Document
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "sample.docx"
            document = Document()
            document.add_paragraph("A short heading")
            document.add_table(rows=1, cols=2)
            document.save(path)
            evidence = extract_docx.extract_docx_evidence(path)
            self.assertEqual(evidence["paragraph_count"], 1)
            self.assertEqual(evidence["table_count"], 1)
            self.assertNotIn(str(path), str(evidence))

    def test_docx_extractor_rejects_macro_container_before_document_parse(self):
        import zipfile
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "macro.docx"
            with zipfile.ZipFile(path, "w") as archive:
                archive.writestr("[Content_Types].xml", "<Types/>")
                archive.writestr("word/vbaProject.bin", b"macro")
            with self.assertRaisesRegex(ValueError, "unsafe_docx_source"):
                extract_docx.extract_docx_evidence(path)


if __name__ == "__main__":
    unittest.main()
